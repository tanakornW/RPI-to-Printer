import requests
import json



r = requests.get(url='https://jsonplaceholder.typicode.com/posts/1')
r = json.dumps(r.json())
loaded_r = json.loads(r)

import barcode
from barcode.writer import ImageWriter
ean = barcode.get('code128', '1234567tanakotnjhy', writer=ImageWriter())
filename = ean.save('testbarcode')
print (ean)

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
a="ธนากร วังทอง 12345"

#changing the page margins
sections = document.sections
for section in sections:
    section.page_height = Cm(5)
    section.page_width = Cm(10)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)


document.add_heading(a,0)
p = document.add_paragraph()
r = p.add_run()
r.add_text(loaded_r['title'])
my_image = document.add_picture('testbarcode.png',height=Cm(1.3)) 
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#document.add_picture('testbarcode.png',height=Cm(1))
document.save('demo.docx')
import os

os.startfile("demo.docx", "print")